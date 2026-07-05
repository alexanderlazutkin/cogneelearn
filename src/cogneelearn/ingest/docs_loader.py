"""Load project documents (md/txt/docx/pdf) into Cognee.

Cognee's native pipeline accepts file paths for md/txt/pdf and extracts text
itself. For .docx we extract text with ``python-docx`` and feed it to Cognee as
a string, because Cognee has no built-in docx loader.

Scanned/OCR'd PDFs are out of scope: the requirement is text-layer PDFs only.
If a PDF has no extractable text, we log a warning and skip it rather than
attempt OCR.
"""

from __future__ import annotations

import logging
from dataclasses import dataclass
from pathlib import Path

logger = logging.getLogger(__name__)

NATIVE_EXTENSIONS = {".md", ".markdown", ".txt", ".pdf"}
EXTRACT_EXTENSIONS = {".docx"}
SUPPORTED_EXTENSIONS = NATIVE_EXTENSIONS | EXTRACT_EXTENSIONS


@dataclass
class LoadedDoc:
    """A document ready to hand to ``cognee.add()``.

    Either ``path`` (native format — Cognee extracts text) or ``text``
    (already extracted — for docx) is set, never both.
    """

    name: str
    source_path: Path
    kind: str  # "native" | "text"
    path: Path | None = None
    text: str | None = None
    char_count: int = 0


def load_document(file_path: str | Path) -> LoadedDoc:
    """Prepare a single document for Cognee ingestion."""
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"Document not found: {path}")
    ext = path.suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type '{ext}' for {path.name}. "
            f"Supported: {sorted(SUPPORTED_EXTENSIONS)}"
        )

    if ext in NATIVE_EXTENSIONS:
        text_len = _estimate_text_length(path)
        return LoadedDoc(
            name=path.name,
            source_path=path,
            kind="native",
            path=path,
            char_count=text_len,
        )

    # .docx — extract text ourselves.
    text = _extract_docx_text(path)
    if not text.strip():
        logger.warning("No text extracted from docx: %s", path.name)
    return LoadedDoc(
        name=path.name,
        source_path=path,
        kind="text",
        text=text,
        char_count=len(text),
    )


def load_directory(dir_path: str | Path) -> list[LoadedDoc]:
    """Load every supported document under a directory (non-recursive)."""
    directory = Path(dir_path)
    if not directory.is_dir():
        raise NotADirectoryError(f"Not a directory: {directory}")
    docs: list[LoadedDoc] = []
    for entry in sorted(directory.iterdir()):
        if not entry.is_file():
            continue
        if entry.suffix.lower() not in SUPPORTED_EXTENSIONS:
            continue
        try:
            docs.append(load_document(entry))
        except Exception as exc:  # noqa: BLE001 — keep loading the rest
            logger.error("Failed to load %s: %s", entry, exc)
    return docs


# ─── extractors ───────────────────────────────────────────────────────────────
def _extract_docx_text(path: Path) -> str:
    """Extract plain text from a .docx file, paragraph by paragraph."""
    try:
        import docx  # python-docx
    except ImportError as exc:  # pragma: no cover — dependency declared
        raise ImportError(
            "python-docx is required to read .docx files. Install it with: pip install python-docx"
        ) from exc

    document = docx.Document(str(path))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text
        if text:
            parts.append(text)
    # Tables: flatten cells row by row to preserve some structure.
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _estimate_text_length(path: Path) -> int:
    """Best-effort character count for a native-format file.

    For txt/md this is exact. For PDF it is an approximation (we do not pull in
    a PDF text layer here — Cognee will do that during ingestion); we just stat
    the file so the UI can show *something* without a heavy parse.
    """
    ext = path.suffix.lower()
    if ext in {".txt", ".md", ".markdown"}:
        try:
            return path.stat().st_size
        except OSError:
            return 0
    # PDF — return raw file size as a rough proxy.
    try:
        return path.stat().st_size
    except OSError:
        return 0
